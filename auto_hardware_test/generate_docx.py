from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

# 切换到目标目录
os.chdir(r'D:\xwechat_files\wxid_uovxpr9gi1t422_bf14\msg\file\2026-03\文档\文档\auto_hardware_test')

# 创建文档
doc = Document()

# 添加标题
title = doc.add_heading('JMeter 压测完整指南', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 添加版本信息
doc.add_paragraph('版本：v1.0          更新日期：2026-04-01')
doc.add_paragraph('适用人群：软件测试工程师、性能测试工程师')
doc.add_paragraph()

# 添加目录说明
doc.add_heading('目录', level=1)
doc.add_paragraph('一、JMeter 安装教程', style='List Bullet')
doc.add_paragraph('二、JMeter 插件安装教程', style='List Bullet')
doc.add_paragraph('三、压测场景配置模板', style='List Bullet')
doc.add_paragraph('四、JMeter 目录结构', style='List Bullet')
doc.add_paragraph('五、常用命令', style='List Bullet')
doc.add_paragraph('六、插件清单', style='List Bullet')
doc.add_paragraph('七、常见问题', style='List Bullet')
doc.add_paragraph()

# 一、JMeter 安装教程
doc.add_heading('一、JMeter 安装教程', level=1)

doc.add_heading('1.1 Windows 安装', level=2)
doc.add_paragraph('1. 检查/安装 Java（JDK 11+）')
doc.add_paragraph('   java -version')
doc.add_paragraph()
doc.add_paragraph('2. 下载 JMeter')
doc.add_paragraph('   访问：https://jmeter.apache.org/download_jmeter.cgi')
doc.add_paragraph('   下载：apache-jmeter-5.6.3.zip')
doc.add_paragraph()
doc.add_paragraph('3. 解压到指定目录')
doc.add_paragraph('   解压到：C:\\Program Files\\apache-jmeter-5.6.3')
doc.add_paragraph()
doc.add_paragraph('4. 配置环境变量（可选）')
doc.add_paragraph('   系统变量 → 新建')
doc.add_paragraph('   变量名：JMETER_HOME')
doc.add_paragraph('   变量值：C:\\Program Files\\apache-jmeter-5.6.3')
doc.add_paragraph('   系统变量 → Path → 新建')
doc.add_paragraph('   %JMETER_HOME%\\bin')
doc.add_paragraph()
doc.add_paragraph('5. 启动 JMeter')
doc.add_paragraph('   双击：C:\\Program Files\\apache-jmeter-5.6.3\\bin\\jmeter.bat')
doc.add_paragraph('   或命令行：jmeter')

doc.add_heading('1.2 Linux/Mac 安装', level=2)
doc.add_paragraph('1. 检查/安装 Java')
doc.add_paragraph('   java -version')
doc.add_paragraph('   如未安装：')
doc.add_paragraph('   sudo apt install openjdk-11-jdk  # Ubuntu/Debian')
doc.add_paragraph('   sudo yum install java-11-openjdk # CentOS/RHEL')
doc.add_paragraph('   brew install openjdk@11          # macOS')
doc.add_paragraph()
doc.add_paragraph('2. 下载 JMeter')
doc.add_paragraph('   wget https://downloads.apache.org//jmeter/binaries/apache-jmeter-5.6.3.tgz')
doc.add_paragraph()
doc.add_paragraph('3. 解压')
doc.add_paragraph('   tar -xzf apache-jmeter-5.6.3.tgz')
doc.add_paragraph('   sudo mv apache-jmeter-5.6.3 /opt/jmeter')
doc.add_paragraph()
doc.add_paragraph('4. 配置环境变量')
doc.add_paragraph('   echo \'export JMETER_HOME=/opt/jmeter\' >> ~/.bashrc')
doc.add_paragraph('   echo \'export PATH=$PATH:$JMETER_HOME/bin\' >> ~/.bashrc')
doc.add_paragraph('   source ~/.bashrc')
doc.add_paragraph()
doc.add_paragraph('5. 启动 JMeter')
doc.add_paragraph('   jmeter')

doc.add_heading('1.3 验证安装', level=2)
doc.add_paragraph('查看版本：')
doc.add_paragraph('   jmeter -v')
doc.add_paragraph('   应该输出：Apache JMeter 5.6.3')

# 二、JMeter 插件安装教程
doc.add_page_break()
doc.add_heading('二、JMeter 插件安装教程', level=1)

doc.add_heading('2.1 插件管理器安装', level=2)
doc.add_paragraph('1. 下载插件管理器')
doc.add_paragraph('   地址：https://jmeter-plugins.org/install/Install/')
doc.add_paragraph('2. 下载文件：JMeterPlugins-Standard-1.4.zip')
doc.add_paragraph('3. 解压后，将 JMeterPlugins-Standard.jar 复制到：')
doc.add_paragraph('   Windows: C:\\Program Files\\apache-jmeter-5.6.3\\lib\\ext\\')
doc.add_paragraph('   Linux:   /opt/jmeter/lib/ext/')
doc.add_paragraph('4. 重启 JMeter')
doc.add_paragraph('5. 在菜单中验证：Options → Plugins Manager（出现则安装成功）')

doc.add_heading('2.2 常用插件推荐', level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '插件名称'
hdr_cells[1].text = '用途'
hdr_cells[2].text = '推荐指数'

table.rows[1].cells[0].text = 'JPG Graphics'
table.rows[1].cells[1].text = '生成更漂亮的性能图表'
table.rows[1].cells[2].text = '⭐⭐⭐⭐⭐'

table.rows[2].cells[0].text = 'PerfMon'
table.rows[2].cells[1].text = '实时监控服务器资源'
table.rows[2].cells[2].text = '⭐⭐⭐⭐⭐'

table.rows[3].cells[0].text = 'Custom Thread Groups'
table.rows[3].cells[1].text = '更灵活的并发控制'
table.rows[3].cells[2].text = '⭐⭐⭐⭐'

table.rows[4].cells[0].text = 'JSON Support'
table.rows[4].cells[1].text = '增强 JSON 处理能力'
table.rows[4].cells[2].text = '⭐⭐⭐⭐'

table.rows[5].cells[0].text = 'WebSocket Samplers'
table.rows[5].cells[1].text = 'WebSocket 协议压测'
table.rows[5].cells[2].text = '⭐⭐⭐'

doc.add_heading('2.3 通过插件管理器安装', level=3)
doc.add_paragraph('1. JMeter GUI → Options → Plugins Manager')
doc.add_paragraph('2. 选择 \"Available Plugins\" 标签')
doc.add_paragraph('3. 推荐安装的插件：')
doc.add_paragraph('   JPG Graphics (Synthetic Blends)')
doc.add_paragraph('   PerfMon (Server Performance Monitor)')
doc.add_paragraph('   Custom Thread Groups')
doc.add_paragraph('   JSON Support')
doc.add_paragraph('   WebSocket Samplers')
doc.add_paragraph('4. 点击 \"Apply Changes and Restart JMeter\"')

# 三、压测场景配置模板
doc.add_page_break()
doc.add_heading('三、压测场景配置模板', level=1)

doc.add_heading('3.1 完整配置模板', level=2)
doc.add_paragraph('='*60)
doc.add_paragraph('           JMeter 压测场景配置模板')
doc.add_paragraph('='*60)
doc.add_paragraph()
doc.add_paragraph('【基本信息】')
doc.add_paragraph('项目名称：______________________________')
doc.add_paragraph('测试场景：______________________________')
doc.add_paragraph('接口协议：HTTP / HTTPS')
doc.add_paragraph()
doc.add_paragraph('【接口信息】')
doc.add_paragraph('接口名称：______________________________')
doc.add_paragraph('请求方法：GET / POST / PUT / DELETE')
doc.add_paragraph('接口地址：______________________________')
doc.add_paragraph()
doc.add_paragraph('【请求配置】')
doc.add_paragraph('1. 请求头（Headers）：')
doc.add_paragraph('   Content-Type: application/json')
doc.add_paragraph('   Authorization: Bearer ${token}')
doc.add_paragraph('   其他：______________________________')
doc.add_paragraph()
doc.add_paragraph('2. 请求参数：')
doc.add_paragraph('   如果是 GET（Query 参数）：参数名 = 参数值')
doc.add_paragraph('   如果是 POST（Body）：JSON 格式或 Form-Data')
doc.add_paragraph()
doc.add_paragraph('【并发配置】')
doc.add_paragraph('虚拟用户数（线程数）：_______')
doc.add_paragraph('准备时长（秒）：_______')
doc.add_paragraph('循环次数：_______')
doc.add_paragraph('持续时间：_______')
doc.add_paragraph()
doc.add_paragraph('【断言配置】')
doc.add_paragraph('预期响应码：200')
doc.add_paragraph('预期响应内容包含：success')
doc.add_paragraph('JSON Path 断言（可选）：$.code = 0')
doc.add_paragraph()
doc.add_paragraph('【参数化配置】')
doc.add_paragraph('需要参数化的字段：____________________')
doc.add_paragraph('CSV 文件名：__________________________')
doc.add_paragraph('CSV 数据示例：value1,value2,value3')
doc.add_paragraph()
doc.add_paragraph('【关联提取】')
doc.add_paragraph('需要从响应中提取的字段：______________')
doc.add_paragraph('提取变量名：__________________________')
doc.add_paragraph('提取表达式（JSON Path）：$.data.token')
doc.add_paragraph()
doc.add_paragraph('【思考时间】')
doc.add_paragraph('固定思考时间：_______毫秒')
doc.add_paragraph('随机思考时间：_______到_______毫秒')

doc.add_heading('3.2 配置示例 - 电商登录接口', level=2)
doc.add_paragraph('【基本信息】')
doc.add_paragraph('项目名称：电商平台性能测试')
doc.add_paragraph('测试场景：用户登录高并发测试')
doc.add_paragraph('接口协议：HTTPS')
doc.add_paragraph()
doc.add_paragraph('【接口信息】')
doc.add_paragraph('接口名称：用户登录')
doc.add_paragraph('请求方法：POST')
doc.add_paragraph('接口地址：https://api.example.com/v1/user/login')
doc.add_paragraph()
doc.add_paragraph('【请求配置】')
doc.add_paragraph('请求头：Content-Type: application/json')
doc.add_paragraph('请求Body（JSON）：')
doc.add_paragraph('{')
doc.add_paragraph('  \"username\": \"${username}\",')
doc.add_paragraph('  \"password\": \"${password}\",')
doc.add_paragraph('  \"captcha\": \"1234\"')
doc.add_paragraph('}')
doc.add_paragraph()
doc.add_paragraph('【并发配置】')
doc.add_paragraph('虚拟用户数：100')
doc.add_paragraph('准备时长：10秒')
doc.add_paragraph('循环次数：10')
doc.add_paragraph()
doc.add_paragraph('【断言配置】')
doc.add_paragraph('预期响应码：200')
doc.add_paragraph('预期响应内容包含：登录成功')
doc.add_paragraph('JSON Path 断言：$.code = 0')
doc.add_paragraph()
doc.add_paragraph('【参数化配置】')
doc.add_paragraph('需要参数化的字段：username, password')
doc.add_paragraph('CSV 文件名：users.csv')
doc.add_paragraph('CSV 数据示例：')
doc.add_paragraph('username,password')
doc.add_paragraph('user001,123456')
doc.add_paragraph('user002,123456')
doc.add_paragraph('user003,123456')
doc.add_paragraph()
doc.add_paragraph('【关联提取】')
doc.add_paragraph('需要提取登录后的 token')
doc.add_paragraph('提取变量名：auth_token')
doc.add_paragraph('提取表达式：$.data.token')

# 四、JMeter 目录结构
doc.add_page_break()
doc.add_heading('四、JMeter 目录结构', level=1)
doc.add_paragraph('apache-jmeter-5.6.3/')
doc.add_paragraph('├── bin/                    # 可执行文件和配置')
doc.add_paragraph('│   ├── jmeter.bat         # Windows 启动脚本')
doc.add_paragraph('│   ├── jmeter.sh          # Linux 启动脚本')
doc.add_paragraph('│   ├── jmeter.properties  # JMeter 主配置')
doc.add_paragraph('│   └── system.properties  # 系统配置')
doc.add_paragraph('├── lib/')
doc.add_paragraph('│   ├── ext/               # 插件 JAR 放这里')
doc.add_paragraph('│   └── junit/')
doc.add_paragraph('├── printable_docs/')
doc.add_paragraph('└── backups/')
doc.add_paragraph()
doc.add_paragraph('.jmx 脚本保存位置（自定义）：')
doc.add_paragraph('Windows: D:\\JMeter_Scripts\\')
doc.add_paragraph('Linux:   ~/jmeter_scripts/')

# 五、常用命令
doc.add_heading('五、常用命令', level=1)
doc.add_paragraph('GUI 模式启动（调试用）：')
doc.add_paragraph('   jmeter')
doc.add_paragraph()
doc.add_paragraph('CLI 模式运行（正式压测）：')
doc.add_paragraph('   jmeter -n -t test.jmx -l result.jtl -e -o report/')
doc.add_paragraph()
doc.add_paragraph('查看版本：')
doc.add_paragraph('   jmeter -v')
doc.add_paragraph()
doc.add_paragraph('远程启动 Server：')
doc.add_paragraph('   ./jmeter-server -Djava.rmi.server.hostname=192.168.1.100')
doc.add_paragraph()
doc.add_paragraph('分布式压测：')
doc.add_paragraph('   jmeter -n -t test.jmx -R 192.168.1.101,192.168.1.102')

doc.add_heading('命令参数说明', level=2)
table2 = doc.add_table(rows=6, cols=2)
table2.style = 'Light Grid Accent 1'
table2.rows[0].cells[0].text = '参数'
table2.rows[0].cells[1].text = '说明'
table2.rows[1].cells[0].text = '-n'
table2.rows[1].cells[1].text = '非 GUI 模式'
table2.rows[2].cells[0].text = '-t'
table2.rows[2].cells[1].text = '指定测试脚本文件'
table2.rows[3].cells[0].text = '-l'
table2.rows[3].cells[1].text = '指定结果日志文件'
table2.rows[4].cells[0].text = '-e'
table2.rows[4].cells[1].text = '测试结束后生成报告'
table2.rows[5].cells[0].text = '-o'
table2.rows[5].cells[1].text = '指定报告输出目录'

# 六、插件清单
doc.add_page_break()
doc.add_heading('六、插件清单', level=1)

doc.add_heading('6.1 推荐安装插件', level=2)
table3 = doc.add_table(rows=6, cols=3)
table3.style = 'Light Grid Accent 1'
table3.rows[0].cells[0].text = '插件'
table3.rows[0].cells[1].text = '功能'
table3.rows[0].cells[2].text = '安装方式'
table3.rows[1].cells[0].text = 'Plugins Manager'
table3.rows[1].cells[1].text = '插件管理器'
table3.rows[1].cells[2].text = '手动安装'
table3.rows[2].cells[0].text = 'PerfMon'
table3.rows[2].cells[1].text = '服务器监控'
table3.rows[2].cells[2].text = 'Plugins Manager'
table3.rows[3].cells[0].text = 'Custom Thread Groups'
table3.rows[3].cells[1].text = '自定义线程组'
table3.rows[3].cells[2].text = 'Plugins Manager'
table3.rows[4].cells[0].text = 'JPG Graphics'
table3.rows[4].cells[1].text = '增强图表'
table3.rows[4].cells[2].text = 'Plugins Manager'
table3.rows[5].cells[0].text = 'WebSocket'
table3.rows[5].cells[1].text = 'WebSocket 协议'
table3.rows[5].cells[2].text = 'Plugins Manager'

doc.add_heading('6.2 PerfMon 服务器监控配置', level=2)
doc.add_paragraph('被测服务器安装 ServerAgent：')
doc.add_paragraph('1. 下载')
doc.add_paragraph('   wget https://github.com/undera/perfmon-agent/releases/download/2.2.3/ServerAgent-2.2.3.zip')
doc.add_paragraph('2. 解压并启动')
doc.add_paragraph('   unzip ServerAgent-2.2.3.zip')
doc.add_paragraph('   cd ServerAgent-2.2.3')
doc.add_paragraph('   ./startAgent.sh --udp-port 4444 --tcp-port 4444')
doc.add_paragraph('3. JMeter 配置')
doc.add_paragraph('   添加 → 监听器 → jp@gc - PerfMon Metrics Collector')
doc.add_paragraph('   Host: 被测服务器IP, Port: 4444')

# 七、常见问题
doc.add_page_break()
doc.add_heading('七、常见问题', level=1)

doc.add_heading('Q1: JMeter 启动报错找不到 Java', level=2)
doc.add_paragraph('解决方法：配置 JAVA_HOME 环境变量')
doc.add_paragraph('Windows: 系统变量 → 新建 → 变量名：JAVA_HOME')
doc.add_paragraph('Linux: export JAVA_HOME=/usr/lib/jvm/java-11')

doc.add_heading('Q2: 插件安装后不显示', level=2)
doc.add_paragraph('1. 确认 JAR 包放在 lib/ext/ 目录')
doc.add_paragraph('2. 重启 JMeter')
doc.add_paragraph('3. 检查 JMeter 版本与插件兼容性')

doc.add_heading('Q3: 压测时内存溢出', level=2)
doc.add_paragraph('修改 jmeter.bat (Windows) 或 jmeter.sh (Linux)：')
doc.add_paragraph('HEAP=\"-Xms1g -Xmx4g -XX:MaxMetaspaceSize=256m\"')

doc.add_heading('Q4: 分布式压测连接失败', level=2)
doc.add_paragraph('1. 检查防火墙，开放端口')
doc.add_paragraph('2. 检查 jmeter.properties 配置')
doc.add_paragraph('   server_port=50000')
doc.add_paragraph('   server.rmi.localport=50000')

# 附录
doc.add_page_break()
doc.add_heading('附录', level=1)

doc.add_heading('A. 下载地址', level=2)
table4 = doc.add_table(rows=4, cols=2)
table4.style = 'Light Grid Accent 1'
table4.rows[0].cells[0].text = '资源'
table4.rows[0].cells[1].text = '地址'
table4.rows[1].cells[0].text = 'JMeter 官网'
table4.rows[1].cells[1].text = 'https://jmeter.apache.org/'
table4.rows[2].cells[0].text = 'JMeter 下载'
table4.rows[2].cells[1].text = 'https://jmeter.apache.org/download_jmeter.cgi'
table4.rows[3].cells[0].text = '插件管理器'
table4.rows[3].cells[1].text = 'https://jmeter-plugins.org/'

doc.add_heading('B. 参考文档', level=2)
doc.add_paragraph('JMeter 官方文档: https://jmeter.apache.org/usermanual/index.html')
doc.add_paragraph('JMeter Plugins Wiki: https://jmeter-plugins.org/wiki/')

doc.add_heading('版本历史', level=1)
table5 = doc.add_table(rows=2, cols=3)
table5.style = 'Light Grid Accent 1'
table5.rows[0].cells[0].text = '版本'
table5.rows[0].cells[1].text = '日期'
table5.rows[0].cells[2].text = '说明'
table5.rows[1].cells[0].text = 'v1.0'
table5.rows[1].cells[1].text = '2026-04-01'
table5.rows[1].cells[2].text = '初始版本'

# 保存文档
doc.save('JMeter压测完整指南.docx')
print('DOCX文件生成成功！')
print('文件位置：D:\\xwechat_files\\wxid_uovxpr9gi1t422_bf14\\msg\\file\\2026-03\\文档\\文档\\auto_hardware_test\\JMeter压测完整指南.docx')
